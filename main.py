import os
import random
import re
import requests
import time
import urllib3
from bs4 import BeautifulSoup
from docx import Document
from loguru import logger


link = input('Ссылка на сайт: ')
start_page = int(input('Стартовая страница загрузки: '))
number_of_pages = int(input('Количество страниц: '))
file_extension = int(input('Расширение файла 1 - txt, 2 - docx: '))
file_extensions = {
    1: 'txt',
    2: 'docx'
}
file_name = re.split(r'/', urllib3.util.parse_url(link).path)[-1]
file_path = f'books/{file_name}.{file_extensions[file_extension]}'
delay = random.randint(20, 30)
bold_paragraph_selector = 'p'

if os.path.exists(file_path) == True:
    os.remove(file_path)
    file = open(file_path, 'w', encoding='utf-8')

if file_extension == 2:
    document = Document()

logger.success('Начало работы скрипта')

for i in range(start_page, number_of_pages + start_page):
    try:
        response = requests.get(f'{link}/?fpage={i}')
        soup = BeautifulSoup(response.text, 'lxml')
        paragraphs = soup.find_all('p')

        if file_extension == 2:
            for paragraph in paragraphs:
                paragraph_classes = paragraph.get('class')
                paragraph_outside_text = paragraph.previous_sibling
                
                try:
                    if isinstance(paragraph_outside_text, str):
                        p = document.add_paragraph(paragraph_outside_text)
                except TypeError:
                    pass

                if bold_paragraph_selector in paragraph_classes:
                    p = document.add_paragraph()
                    p.add_run(paragraph.text).bold = True
                else:
                    p = document.add_paragraph(paragraph.text)

            logger.success(f'Данные страницы {i} загружены')
        else:
            with open(file_path, 'a', encoding='utf-8') as file:
                for paragraph in paragraphs:
                    paragraph_outside_text = paragraph.previous_sibling
                
                    if paragraph_outside_text != None and len(paragraph_outside_text) > 0:
                        file.write(f'{paragraph_outside_text.text}\n')

                    file.write(f'{paragraph.text}\n')

            logger.success(f'Данные страницы {i} добавлены в файл')
        time.sleep(delay)

    except Exception:
        logger.error(f'Не удалось загрузить или записать данные страницы {i}')

if file_extension == 2:
    document.save(file_path)
    logger.success(f'Данные добавлены в файл {file_path}')

logger.success('Конец работы скрипта')